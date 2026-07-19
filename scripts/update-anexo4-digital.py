"""Actualiza Anexo 4 / metodología: instrumentos digitales y protocolo Instrumento 4."""
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

path = (
    r"c:\Users\geova\OneDrive\Desktop\10 semestre\Proyecto de graduación II"
    r"\Version editada PG2\Proyecto de graduación II.docx"
)
doc = Document(path)


def set_para(i: int, text: str) -> None:
    p = doc.paragraphs[i]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.text = text


set_para(
    405,
    "Instrumentos de línea base. Para la medición inicial se aplicarán instrumentos digitales, "
    "con el fin de facilitar el registro, la tabulación y el análisis de resultados, evitando el uso de formatos manuscritos. "
    "En concreto: (a) la ficha de observación y registro de tiempos (Instrumento 1) se aplicará en una hoja de cálculo digital "
    "(Google Sheets o Microsoft Excel), donde cada fila corresponderá a una cotización observada y permitirá calcular automáticamente "
    "los promedios de la sección C; (b) los cuestionarios al personal interno y a clientes (Instrumentos 2 y 3) se aplicarán mediante "
    "Google Forms, con exportación a hoja de cálculo para el análisis; y (c) la guía de observación del cliente (Instrumento 4) se aplicará "
    "también en formato digital (Google Forms o Google Sheets), registrando si ocurre cada aspecto del proceso, el tiempo estimado y observaciones. "
    "Estos instrumentos pertenecen a la medición investigativa (línea base / post-implementación) y no sustituyen el tablero operativo del sistema; "
    "sus resultados se presentan en el informe académico y, de forma opcional, en un panel simple de resultados del estudio (por ejemplo Looker Studio "
    "conectado a Sheets). Con ello se obtendrá evidencia numérica y perceptual del estado inicial del negocio, necesaria para contrastar los resultados "
    "obtenidos tras la puesta en marcha del sistema propuesto (Anexo 4).",
)

set_para(
    846,
    "Los siguientes instrumentos constituyen la medición inicial o línea base del proceso comercial vigente de Ferromaderas, "
    "previa a la implementación de la plataforma web. Su aplicación permite cuantificar el estado actual del negocio y contrastarlo "
    "posteriormente con los resultados obtenidos tras la puesta en marcha del sistema propuesto, en coherencia con la matriz operacional (Anexo 3). "
    "Todos los instrumentos se aplican en formato digital: Instrumento 1 en Google Sheets (o Excel); Instrumentos 2 y 3 en Google Forms; "
    "Instrumento 4 en Google Forms o Google Sheets. Los resultados se documentan con enlace al formulario/hoja, captura de pantalla y tabla resumen "
    "en el informe. No se incorporan como módulo del panel administrativo de Ferromaderas, porque miden el estudio (antes/después), no la operación diaria del sistema.",
)

set_para(847, "Instrumento 1. Ficha de observación y registro de tiempo (Google Sheets / Excel)")
set_para(859, "Instrumento 2. Cuestionario Google Forms — Personal interno")
set_para(861, "Instrumento 3. Cuestionario Google Forms — Clientes")
set_para(863, "Instrumento 4. Guía de observación del cliente (Google Forms o Google Sheets)")

protocol = (
    "Protocolo de medición del Instrumento 4. Se observa el proceso real de atención (presencial o WhatsApp) sin intervenir. "
    "Por cada observación se registra: (1) si ocurre cada aspecto (Sí/No); (2) el tiempo aproximado en minutos desde la solicitud "
    "hasta la entrega de la cotización, cuando aplique; y (3) observaciones cualitativas. La muestra mínima sugerida es de 8 a 12 "
    "observaciones en días hábiles distintos. Indicadores derivados: porcentaje de casos en que se abre Di-Chara; tiempo promedio "
    "de espera del cliente; porcentaje de casos con seguimiento registrado. Estos valores alimentan la línea base y se comparan, "
    "en la medición post, con el uso de la plataforma web (cotización en línea, seguimiento digital y menor dependencia de Di-Chara "
    "en la primera respuesta al cliente)."
)

p864 = doc.paragraphs[864]
if not (p864.text or "").strip():
    set_para(864, protocol)
else:
    ref = doc.paragraphs[863]._p
    new_el = OxmlElement("w:p")
    ref.addnext(new_el)
    Paragraph(new_el, doc.paragraphs[863]._parent).text = protocol

doc.save(path)
print("OK: Anexo 4 y metodología actualizados")
for i in [405, 846, 847, 859, 861, 863, 864]:
    t = Document(path).paragraphs[i].text
    print(f"[{i}] {t[:160]}...")
